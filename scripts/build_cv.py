import re
import os
import tempfile
from datetime import datetime, timezone
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any
from collections.abc import Sequence

DEFAULT_SOURCE_DATE_EPOCH = 1785240000  # 2026-07-28 12:00:00 UTC
BUILD_EPOCH = int(os.environ.get("SOURCE_DATE_EPOCH", DEFAULT_SOURCE_DATE_EPOCH))
os.environ.setdefault("SOURCE_DATE_EPOCH", str(BUILD_EPOCH))
ZIP_DATE_TIME = datetime.fromtimestamp(max(BUILD_EPOCH, 315532800), timezone.utc).timetuple()[:6]

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Mm, Pt
from reportlab import rl_config
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import KeepTogether, PageBreak, Paragraph, SimpleDocTemplate

rl_config.invariant = 1

ROOT = Path(__file__).resolve().parents[1]
PUBLIC = ROOT / "public"
PUBLIC.mkdir(parents=True, exist_ok=True)
DOCX_PATH = PUBLIC / "Joshua_Nwachinemere_CV.docx"
PDF_PATH = PUBLIC / "Joshua_Nwachinemere_CV.pdf"

NAME = "JOSHUA NWACHINEMERE"
HEADLINE = "AI Engineer | Python, Retrieval Systems, Multimodal AI & ML Evaluation"
EMAIL = "joshua0nwachinemere@gmail.com"
PHONE = "+234 913 148 0096"

GITHUB_URL = "https://github.com/dk3yyyy"
LINKEDIN_URL = "https://www.linkedin.com/in/joshua-nwachinemere"
PORTFOLIO_URL = "https://joshua-nwachinemere.pages.dev"
URL_PATTERN = re.compile(r"https://[^\s|]+")

SUMMARY = (
    "AI Engineer building Python systems for retrieval, multimodal inference, speech, model APIs, and ML evaluation. "
    "Built local review intelligence with citation validation and benchmarked evaluation, a privacy-oriented macOS "
    "assistant, and a temporal forecasting pipeline. Contributed eight merged fixes and tests across AI and data "
    "infrastructure projects, including OpenAI Agents SDK and Pydantic AI Harness."
)

EXPERIENCE = [
    {
        "role": "AI Engineer",
        "org": "VolyxAI | Independent AI product development",
        "date": "Nov 2025 - Present",
        "bullets": [
            "Build applied AI systems integrating retrieval, multimodal inputs, voice workflows, model APIs, and structured outputs.",
            "Develop Python services and n8n workflows with webhooks, validation, retries, logging, human approval, and provider fallback.",
        ],
        "tools": "Python, FastAPI, n8n, REST APIs, webhooks, Azure AI Foundry, OpenAI-compatible APIs, PostgreSQL, Redis",
    },
    {
        "role": "Python Automation Developer",
        "org": "Freelance | Independent paid client work",
        "date": "Jan 2023 - Present",
        "bullets": [
            "Developed streaming utilities for large text datasets, including domain filtering, deduplication, splitting, and list curation.",
            "Built Telegram integrations and asynchronous services using PostgreSQL, Redis, and Docker; integrated external APIs with explicit error handling, rate limits, and operational logging.",
            "Developed licensing and key-management components for proprietary Python tools, and published reusable command-line utilities for splitting and sorting large text datasets.",
        ],
        "tools": "Python, scripting, asyncio, Telegram Bot API, aiohttp, PostgreSQL, Redis, Docker, REST APIs",
    },
]

OPEN_SOURCE_CONTRIBUTIONS = [
    {
        "name": "OpenAI Agents Python SDK | Merged PR #3991",
        "stack": "Python, WebSockets, retry policies, pytest | Jul 2026",
        "bullets": ["Made transient pre-response WebSocket server errors follow the SDK's retry policy while preserving non-transient handling and replay-safety checks. PR: https://github.com/openai/openai-agents-python/pull/3991"],
    },
    {
        "name": "Pydantic AI Harness | Merged PR #503",
        "stack": "Python, agent recovery, Code Mode, pytest | Jul 2026",
        "bullets": ["Surfaced unexpected Code Mode execution failures to the model as retries after resetting invalid session state, enabling recovery from lost imports and variables. PR: https://github.com/pydantic/pydantic-ai-harness/pull/503"],
    },
    {
        "name": "Mellea | Merged PR #1471",
        "stack": "Python, OpenTelemetry tracing, async tests, pytest | Jul 2026",
        "bullets": ["Added deterministic mocked-backend tests for async span timing, context propagation, token usage, span lifetime, and consecutive generations. PR: https://github.com/generative-computing/mellea/pull/1471"],
    },
    {
        "name": "FastStream | Merged PR #2961",
        "stack": "Python, FastAPI compatibility, dependency injection, pytest | Jul 2026",
        "bullets": ["Restored FastAPI 0.140 compatibility for slotted Dependant objects while retaining native dependency fields and FastStream schema metadata. PR: https://github.com/ag2ai/faststream/pull/2961"],
    },
    {
        "name": "Additional verified upstream merges | 4 of 8 merged PRs",
        "stack": "Correctness, dependency errors, schema generation, workflow scoping | Jul 2026",
        "bullets": [
            "Apache Arrow Rust: https://github.com/apache/arrow-rs/pull/10486 | Altair: https://github.com/vega/altair/pull/4089 | FastStream FastAPI: https://github.com/faststream-community/faststream_fastapi/pull/2 | Calkit: https://github.com/calkit/calkit/pull/1028",
        ],
    },
]

PROJECTS = [
    {
        "name": "Local Review Intelligence | Local Retrieval & Evaluation",
        "stack": "Python, Ollama, ChromaDB, LangChain, Streamlit",
        "bullets": [
            "Built a local-first review intelligence application with adaptive CSV mapping, content-addressed Chroma storage, reconciled indexing, a Streamlit dashboard, and an installable CLI.",
            "Implemented stable source IDs, validated citations, bounded answer repair, safe abstention, and a versioned 30-case evaluation comparing semantic retrieval with BM25.",
            "Source: https://github.com/dk3yyyy/local_AI_agent",
        ],
    },
    {
        "name": "Volyx Lens | Privacy-First Context-Aware AI Assistant",
        "stack": "JavaScript, Electron, Swift, Azure AI Foundry, Multimodal APIs",
        "bullets": [
            "Built and hardened a privacy-oriented macOS assistant combining selected screen context, microphone input, meeting audio, local OCR, relevance-ranked retrieval, and multiple AI providers.",
            "Implemented bounded context memory, provider routing, explicit data-sharing controls, sandboxed Electron boundaries, automated tests, secret scanning, and release checks.",
            "Source: https://github.com/dk3yyyy/volyx-lens",
        ],
    },
    {
        "name": "Football Forecasting Lab | Temporal ML Evaluation Pipeline",
        "stack": "Python, XGBoost, scikit-learn, FastAPI, Streamlit, SQLAlchemy",
        "bullets": [
            "Built data scrapers, temporal feature engineering, XGBoost outcome and Poisson goal models, a FastAPI service, and a Streamlit dashboard.",
            "Evaluated with rolling-origin testing across 1,140 Premier League matches: 53.77% outcome accuracy versus a 56.70% bookmaker-implied benchmark; the model did not outperform the benchmark.",
            "Source: https://github.com/dk3yyyy/football_predictor",
        ],
    },
]

SKILLS = [
    "Languages: Python, SQL",
    "AI Engineering: retrieval, multimodal inference, voice AI, structured outputs, provider routing, model evaluation",
    "Models & Providers: Azure AI Foundry, OpenAI-compatible APIs, Deepgram",
    "ML & Data: scikit-learn, XGBoost, pandas, feature engineering, temporal evaluation",
    "Backend & Automation: FastAPI, asyncio, REST APIs, n8n, webhooks, SQLAlchemy",
    "Data & Infrastructure: PostgreSQL, SQLite, Redis, Docker, GitHub Actions, Linux",
]

CERTIFICATIONS = [
    {"text": "Scientific Computing with Python Developer Certification | freeCodeCamp | May 2026", "url": "https://www.freecodecamp.org/certification/joshua_nwachinemere/scientific-computing-with-python-v7"},
    {"text": "Google AI Professional Certificate | Google/Coursera | February 2026", "url": "https://www.coursera.org/account/accomplishments/professional-cert/L1UIFMPUME30"},
    {"text": "Model Context Protocol: Advanced Topics | Anthropic training | March 2026", "url": "https://verify.skilljar.com/c/fwqra86yief7"},
]

EDUCATION = [
    "MSc Artificial Intelligence · Northumbria University · September 2026 intake",
    "Bachelor of Technology (BTech), Mathematics · Federal University of Technology, Owerri (FUTO) · 2016–2021",
]


def configure_docx() -> Any:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    section.different_first_page_header_footer = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.add_run("Joshua Nwachinemere | AI Engineer").font.size = Pt(8)
    doc.core_properties.title = "Joshua Nwachinemere CV"
    doc.core_properties.author = "Joshua Nwachinemere"
    doc.core_properties.subject = "AI Engineer CV"
    created = modified = datetime.fromtimestamp(BUILD_EPOCH, timezone.utc)
    doc.core_properties.created = created
    doc.core_properties.modified = modified
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    return doc


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.rels._next_rId
    paragraph.part.rels.add_relationship(RT.HYPERLINK, url, relationship_id, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_text_with_links(paragraph, text: str) -> None:
    cursor = 0
    for match in URL_PATTERN.finditer(text):
        paragraph.add_run(text[cursor:match.start()])
        url = match.group(0)
        add_hyperlink(paragraph, url, url)
        cursor = match.end()
    paragraph.add_run(text[cursor:])


def add_docx_contact(doc: Document) -> None:
    first = doc.add_paragraph()
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first.paragraph_format.space_after = Pt(0)
    add_hyperlink(first, EMAIL, f"mailto:{EMAIL}")
    first.add_run(f" | {PHONE}")

    second = doc.add_paragraph()
    second.alignment = WD_ALIGN_PARAGRAPH.CENTER
    second.paragraph_format.space_after = Pt(4)
    add_hyperlink(second, "github.com/dk3yyyy", GITHUB_URL)
    second.add_run(" | ")
    add_hyperlink(second, "linkedin.com/in/joshua-nwachinemere", LINKEDIN_URL)
    second.add_run(" | ")
    add_hyperlink(second, "joshua-nwachinemere.pages.dev", PORTFOLIO_URL)


def add_docx_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)


def add_docx_bullets(doc: Any, bullets: Sequence[str | dict[str, str]]) -> None:
    numbering = doc.part.numbering_part.element
    if not hasattr(doc, "_cv_bullet_num_id"):
        abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
        num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
        abstract_id = max(abstract_ids, default=-1) + 1
        num_id = max(num_ids, default=0) + 1
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        level = OxmlElement("w:lvl"); level.set(qn("w:ilvl"), "0")
        fmt = OxmlElement("w:numFmt"); fmt.set(qn("w:val"), "bullet")
        text = OxmlElement("w:lvlText"); text.set(qn("w:val"), "•")
        level.extend([fmt, text]); abstract.append(level); numbering.append(abstract)
        num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
        ref = OxmlElement("w:abstractNumId"); ref.set(qn("w:val"), str(abstract_id)); num.append(ref); numbering.append(num)
        doc._cv_bullet_num_id = num_id
    for bullet in bullets:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(0)
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId"); num_id.set(qn("w:val"), str(doc._cv_bullet_num_id))
        num_pr.extend([ilvl, num_id]); p._p.get_or_add_pPr().append(num_pr)
        if isinstance(bullet, dict):
            p.add_run(f"{bullet['text']} | ")
            add_hyperlink(p, "Verify credential", bullet["url"])
        else:
            add_text_with_links(p, bullet)


def update_extended_properties(path: Path) -> None:
    with zipfile.ZipFile(path) as source:
        source_info = {info.filename: info for info in source.infolist()}
        parts = {name: source.read(name) for name in source.namelist()}
    root = ET.fromstring(parts["docProps/app.xml"])
    ns = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
    for tag in ("Pages", "Lines", "Characters", "CharactersWithSpaces"):
        for node in root.findall(f"{{{ns}}}{tag}"):
            root.remove(node)
    document_root = ET.fromstring(parts["word/document.xml"])
    paragraphs = len(document_root.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"))
    words = sum(len((node.text or "").split()) for node in document_root.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"))
    for tag, value in (("Paragraphs", paragraphs), ("Words", words)):
        node = root.find(f"{{{ns}}}{tag}")
        if node is None:
            node = ET.SubElement(root, f"{{{ns}}}{tag}")
        node.text = str(value)
    parts["docProps/app.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    fd, temp = tempfile.mkstemp(suffix=".docx", dir=path.parent); os.close(fd)
    try:
        with zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as target:
            for name, data in parts.items():
                original = source_info[name]
                info = zipfile.ZipInfo(name, ZIP_DATE_TIME)
                info.compress_type = zipfile.ZIP_DEFLATED
                info.create_system = original.create_system
                info.external_attr = original.external_attr
                info.internal_attr = original.internal_attr
                target.writestr(info, data)
        os.replace(temp, path)
    finally:
        if os.path.exists(temp): os.unlink(temp)


def build_docx(output_path: Path = DOCX_PATH) -> None:
    doc = configure_docx()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(NAME)
    r.bold = True
    r.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(HEADLINE)
    r.bold = True
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(1)

    add_docx_contact(doc)

    add_docx_heading(doc, "PROFILE")
    p = doc.add_paragraph(SUMMARY)
    p.paragraph_format.space_after = Pt(2)

    add_docx_heading(doc, "SELECTED AI & ML PROJECTS")
    for project in PROJECTS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(project["name"])
        r.bold = True
        p = doc.add_paragraph(project["stack"])
        p.paragraph_format.space_after = Pt(0)
        p.runs[0].italic = True
        add_docx_bullets(doc, project["bullets"])

    add_docx_heading(doc, "OPEN-SOURCE CONTRIBUTIONS")
    for index, contribution in enumerate(OPEN_SOURCE_CONTRIBUTIONS):
        if index == 2:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(contribution["name"])
        r.bold = True
        p = doc.add_paragraph(contribution["stack"])
        p.paragraph_format.space_after = Pt(0)
        p.runs[0].italic = True
        add_docx_bullets(doc, contribution["bullets"])

    add_docx_heading(doc, "EXPERIENCE")
    for item in EXPERIENCE:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(item["role"])
        r.bold = True
        p.add_run(f" | {item['date']}").italic = True
        p = doc.add_paragraph(item["org"])
        p.paragraph_format.space_after = Pt(0)
        add_docx_bullets(doc, item["bullets"])
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("Key tools: ")
        r.bold = True
        p.add_run(item["tools"])

    add_docx_heading(doc, "CORE SKILLS")
    add_docx_bullets(doc, SKILLS)

    add_docx_heading(doc, "EDUCATION")
    for line in EDUCATION:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)

    add_docx_heading(doc, "CERTIFICATIONS & TRAINING")
    add_docx_bullets(doc, CERTIFICATIONS)

    doc.save(output_path)
    update_extended_properties(output_path)


def pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "name": ParagraphStyle("CVName", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=16, leading=18, alignment=TA_CENTER, spaceAfter=2),
        "headline": ParagraphStyle("CVHeadline", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=10.5, leading=12, alignment=TA_CENTER, spaceAfter=2),
        "contact": ParagraphStyle("CVContact", parent=styles["Normal"], fontName="Helvetica", fontSize=8.1, leading=10, alignment=TA_CENTER, spaceAfter=5),
        "section": ParagraphStyle("CVSection", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=colors.HexColor("#111111"), spaceBefore=10, spaceAfter=5, borderWidth=0, borderPadding=0),
        "closing_section": ParagraphStyle("CVClosingSection", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=colors.HexColor("#111111"), spaceBefore=5, spaceAfter=3, borderWidth=0, borderPadding=0),
        "role": ParagraphStyle("CVRole", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=10.8, leading=13.5, spaceBefore=5, spaceAfter=1.5),
        "org": ParagraphStyle("CVOrg", parent=styles["Normal"], fontName="Helvetica-Oblique", fontSize=10, leading=12.5, spaceAfter=1.5),
        "body": ParagraphStyle("CVBody", parent=styles["Normal"], fontName="Helvetica", fontSize=10.7, leading=13.7, spaceAfter=2.5),
        "tools": ParagraphStyle("CVTools", parent=styles["Normal"], fontName="Helvetica", fontSize=9.9, leading=12.5, spaceAfter=5),
        "contribution_role": ParagraphStyle("CVContributionRole", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=10, leading=11.5, spaceBefore=3, spaceAfter=.5),
        "contribution_body": ParagraphStyle("CVContributionBody", parent=styles["Normal"], fontName="Helvetica", fontSize=10.2, leading=12, spaceAfter=1),
        "credential": ParagraphStyle("CVCredential", parent=styles["Normal"], fontName="Helvetica", fontSize=10, leading=11.5, spaceAfter=0),
    }


def bullet_list(lines: Sequence[str | dict[str, str]], style: ParagraphStyle) -> KeepTogether:
    bullet_style = ParagraphStyle(
        f"{style.name}Bullet",
        parent=style,
        leftIndent=10,
        firstLineIndent=-7,
        spaceAfter=1,
    )
    return KeepTogether(
        [Paragraph(
            "- " + (
                f"{line['text']} | <link href=\"{line['url']}\" color=\"#1155CC\">Verify credential</link>"
                if isinstance(line, dict)
                else URL_PATTERN.sub(
                    lambda match: f'<link href="{match.group(0)}" color="#1155CC">{match.group(0)}</link>',
                    line,
                )
            ),
            bullet_style,
        ) for line in lines]
    )


def draw_later_page_header(canvas, document) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#444444"))
    canvas.drawRightString(A4[0] - 14 * mm, A4[1] - 8 * mm, "Joshua Nwachinemere | AI Engineer")
    canvas.setStrokeColor(colors.HexColor("#CCCCCC"))
    canvas.line(14 * mm, A4[1] - 10 * mm, A4[0] - 14 * mm, A4[1] - 10 * mm)
    canvas.restoreState()


def build_pdf(output_path: Path = PDF_PATH) -> None:
    styles = pdf_styles()
    contact_markup = (
        f'<link href="mailto:{EMAIL}" color="#1155CC">{EMAIL}</link> | {PHONE}<br/>'
        f'<link href="{GITHUB_URL}" color="#1155CC">github.com/dk3yyyy</link> | '
        f'<link href="{LINKEDIN_URL}" color="#1155CC">linkedin.com/in/joshua-nwachinemere</link> | '
        f'<link href="{PORTFOLIO_URL}" color="#1155CC">joshua-nwachinemere.pages.dev</link>'
    )
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=14 * mm,
        rightMargin=14 * mm,
        topMargin=14 * mm,
        bottomMargin=10 * mm,
        title="Joshua Nwachinemere CV",
        author="Joshua Nwachinemere",
    )
    story: list[Any] = [
        Paragraph(NAME, styles["name"]),
        Paragraph(HEADLINE, styles["headline"]),
        Paragraph(contact_markup, styles["contact"]),
        Paragraph("PROFILE", styles["section"]),
        Paragraph(SUMMARY, styles["body"]),
        Paragraph("SELECTED AI & ML PROJECTS", styles["section"]),
    ]
    for project in PROJECTS:
        story.append(Paragraph(project["name"], styles["role"]))
        story.append(Paragraph(project["stack"], styles["org"]))
        story.append(bullet_list(project["bullets"], styles["body"]))

    story.append(Paragraph("OPEN-SOURCE CONTRIBUTIONS", styles["section"]))
    for index, contribution in enumerate(OPEN_SOURCE_CONTRIBUTIONS):
        if index == 2:
            story.append(PageBreak())
        story.append(Paragraph(contribution["name"], styles["contribution_role"]))
        story.append(Paragraph(contribution["stack"], styles["org"]))
        story.append(bullet_list(contribution["bullets"], styles["contribution_body"]))

    story.append(Paragraph("EXPERIENCE", styles["section"]))
    for item in EXPERIENCE:
        story.append(Paragraph(f"{item['role']} | {item['date']}", styles["role"]))
        story.append(Paragraph(item["org"], styles["org"]))
        story.append(bullet_list(item["bullets"], styles["body"]))
        story.append(Paragraph(f"<b>Key tools:</b> {item['tools']}", styles["tools"]))

    story.append(Paragraph("CORE SKILLS", styles["section"]))
    story.append(bullet_list(SKILLS, styles["body"]))

    story.append(Paragraph("EDUCATION", styles["closing_section"]))
    for line in EDUCATION:
        story.append(Paragraph(line, styles["credential"]))

    story.append(Paragraph("CERTIFICATIONS & TRAINING", styles["closing_section"]))
    story.append(bullet_list(CERTIFICATIONS, styles["credential"]))

    doc.build(story, onLaterPages=draw_later_page_header)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
